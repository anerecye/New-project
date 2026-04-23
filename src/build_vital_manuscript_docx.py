from __future__ import annotations

import csv
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parents[1]
MANUSCRIPT_MD = BASE_DIR / "manuscript_vital_updated.md"
OUTPUT_DOCX = BASE_DIR / "manuscript_vital_with_figures_tables.docx"

# Main-text figures are inserted explicitly with Markdown image tags in
# manuscript_vital_updated.md. Older long-form appendix figure insertions remain
# reproducible from repository outputs but are no longer injected into the main
# Word manuscript.
FIGURE_INSERTIONS = {}

APPENDIX_TABLES = [
    (
        "Appendix Table A1. Baseline comparison against VITAL red",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_method_comparison.csv",
        [
            "method",
            "true_positives",
            "false_positives",
            "false_negatives",
            "true_negatives",
            "precision",
            "recall",
            "specificity",
        ],
    ),
    (
        "Appendix Table A2. VITAL threshold sweep",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_threshold_sweep.csv",
        [
            "score_threshold",
            "flagged_count",
            "true_positives",
            "false_positives",
            "false_negatives",
            "precision",
            "recall",
            "specificity",
        ],
    ),
    (
        "Appendix Table A3. Historical enrichment from 2023 to current ClinVar",
        BASE_DIR / "data" / "processed" / "arrhythmia_2023_01_to_current_vital_historical_enrichment.csv",
        [
            "scope",
            "target",
            "baseline_variant_count",
            "event_count",
            "overall_event_rate",
            "vital_red_count",
            "vital_red_event_count",
            "vital_red_event_rate",
            "red_enrichment_vs_overall",
        ],
    ),
    (
        "Appendix Table A4. Current AC threshold sensitivity and red-set stability",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_ac_threshold_sensitivity.csv",
        [
            "ac_threshold",
            "method",
            "ac_supported_frequency_flag_count_all_observed",
            "vital_red_count_all_observed",
            "vital_red_variants",
            "vital_red_gained_vs_ac_ge_20_variants",
            "vital_red_lost_vs_ac_ge_20_variants",
            "false_positives",
            "precision",
            "precision_ci_low",
            "precision_ci_high",
        ],
    ),
    (
        "Appendix Table A5. Historical AC threshold sensitivity and red-set stability",
        BASE_DIR / "data" / "processed" / "arrhythmia_2023_01_to_current_vital_historical_ac_threshold_sensitivity.csv",
        [
            "scope",
            "target",
            "ac_threshold",
            "method",
            "ac_supported_frequency_flag_count_scope",
            "vital_red_count_at_ac_threshold",
            "vital_red_variants",
            "vital_red_gained_vs_ac_ge_20_variants",
            "vital_red_lost_vs_ac_ge_20_variants",
            "true_positives",
            "false_positives",
            "precision",
            "precision_ci_low",
            "precision_ci_high",
        ],
    ),
    (
        "Appendix Table A6. Review fragility summary",
        BASE_DIR / "data" / "processed" / "arrhythmia_review_fragility_summary.csv",
        [
            "review_strength",
            "variant_count",
            "frequency_flag_count",
            "ac_supported_frequency_count",
            "vital_red_count",
            "share_of_ac_supported_frequency_variants",
            "share_of_vital_red_variants",
        ],
    ),
    (
        "Appendix Table A7. VITAL expert-weight sensitivity by red-priority variant",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_weight_sensitivity_variant_profile_table.csv",
        [
            "gene",
            "clinvar_id",
            "weight_profile",
            "sensitivity_vital_score",
            "retained_red_under_profile",
            "red_set_change",
            "stability_class",
        ],
    ),
    (
        "Appendix Table A7b. One-weight-at-a-time VITAL sensitivity",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_one_weight_sensitivity.csv",
        [
            "component_varied",
            "weight_multiplier",
            "tested_weight",
            "red_priority_count",
            "red_priority_variants",
            "sensitivity_note",
        ],
    ),
    (
        "Appendix Table A8. Historical threshold calibration best thresholds",
        BASE_DIR / "data" / "processed" / "arrhythmia_2023_01_to_current_vital_historical_threshold_calibration_best.csv",
        [
            "target",
            "calibration_mode",
            "best_threshold_by_enrichment",
            "best_threshold_flagged_count",
            "best_threshold_flagged_event_count",
            "best_threshold_enrichment_vs_overall",
            "threshold_70_flagged_count",
            "threshold_70_flagged_event_count",
            "threshold_70_enrichment_vs_overall",
            "threshold_70_matches_best",
        ],
    ),
    (
        "Appendix Table A9. Historical calibration field-mapping fix",
        BASE_DIR / "data" / "processed" / "arrhythmia_2023_01_to_current_vital_historical_threshold_calibration_mapping_fix_summary.csv",
        [
            "target",
            "calibration_mode",
            "before_fix_threshold_70_flagged_count",
            "before_fix_threshold_70_flagged_event_count",
            "after_fix_threshold_70_flagged_count",
            "after_fix_threshold_70_flagged_event_count",
            "after_fix_threshold_70_enrichment_vs_overall",
        ],
    ),
    (
        "Appendix Table A10. VITAL pathogenicity tension continuum by score band",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_frequency_function_discordance_summary.csv",
        [
            "vital_score_band_20pt",
            "band_variant_count",
            "weak_or_single_review_percent",
            "popmax_af_gt_1e_4_percent",
            "ac_supported_frequency_percent",
            "indel_or_duplication_percent",
            "canonical_or_atypical_function_percent",
            "median_max_frequency_signal",
            "median_qualifying_ac",
        ],
    ),
    (
        "Appendix Table A11. VITAL signal reorganization layers",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_signal_reorganization_summary.csv",
        [
            "signal_layer",
            "variant_count",
            "median_vital_score",
            "weak_or_single_review_percent",
            "popmax_af_gt_1e_4_percent",
            "ac_supported_frequency_percent",
            "indel_or_duplication_percent",
            "canonical_or_atypical_function_percent",
            "median_max_frequency_signal",
            "median_qualifying_ac",
        ],
    ),
    (
        "Appendix Table A12. LOF subtype discordance across the VITAL continuum",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_lof_subtype_discordance_summary.csv",
        [
            "lof_subtype",
            "lof_variant_count",
            "naive_af_flag_count",
            "naive_af_flag_percent",
            "ac_supported_frequency_count",
            "ac_supported_frequency_percent",
            "vital_60_100_count",
            "vital_60_100_percent",
            "vital_red_count",
            "max_vital_score",
            "median_max_frequency_signal",
            "max_frequency_signal",
            "max_qualifying_ac",
        ],
    ),
    (
        "Appendix Table A13. Live manual-review snapshot for VITAL-red variants",
        BASE_DIR / "data" / "processed" / "vital_red_manual_review_live_check.csv",
        [
            "gene",
            "clinvar_id",
            "live_check_date",
            "live_clinvar_status",
            "condition",
            "review_support",
            "submitter_conflict_status",
            "clinvar_publications_or_mentions",
            "why_suspicious_but_not_resolved",
        ],
    ),
    (
        "Appendix Table A14. Independent 3,000-variant current cross-disease validation",
        BASE_DIR / "data" / "processed" / "vital_cross_disease_3000_cross_disease_validation_summary.csv",
        [
            "scope",
            "metric",
            "count",
            "denominator",
            "rate",
            "rate_ci_low",
            "rate_ci_high",
            "expected_per_10000",
            "expected_per_10000_ci_low",
            "expected_per_10000_ci_high",
        ],
    ),
    (
        "Appendix Table A15. Independent 2023-to-current cross-disease historical enrichment",
        BASE_DIR / "data" / "processed" / "vital_cross_disease_3000_2023_01_to_current_vital_historical_enrichment.csv",
        [
            "scope",
            "target",
            "baseline_variant_count",
            "event_count",
            "overall_event_rate",
            "vital_red_count",
            "vital_red_event_count",
            "vital_red_event_rate",
            "red_enrichment_vs_overall",
            "red_enrichment_vs_nonred",
            "red_enrichment_vs_nonred_ci_low",
            "red_enrichment_vs_nonred_ci_high",
        ],
    ),
    (
        "Appendix Table A16. VITAL-red biological contrast case studies",
        BASE_DIR / "data" / "processed" / "arrhythmia_vital_biological_contrast_cases.csv",
        [
            "variant",
            "vital",
            "af_signal",
            "loeuf_constraint_argument",
            "heart_expression",
            "expected_mechanism",
            "reality_interpretation",
        ],
    ),
    (
        "Appendix Table A17. Pre-specified blinded expert validation protocol",
        BASE_DIR / "data" / "processed" / "vital_independent_validation_protocol.csv",
        [
            "protocol_element",
            "pre_specified_value",
            "rationale",
        ],
    ),
    (
        "Appendix Table A18. AC reliability strata for frequency-flagged variants",
        BASE_DIR / "data" / "processed" / "vital_ac_reliability_strata.csv",
        [
            "ac_reliability_stratum",
            "frequency_flagged_variants",
            "vital_red_variants",
            "red_variant_ids",
            "interpretation",
        ],
    ),
    (
        "Appendix Table A19. Extended AC threshold sensitivity",
        BASE_DIR / "data" / "processed" / "vital_ac_threshold_sensitivity_extended.csv",
        [
            "ac_threshold",
            "red_queue_count",
            "red_queue_variants",
            "composition_note",
        ],
    ),
    (
        "Appendix Table A20. Required detectability fields before gray-to-actionable routing",
        BASE_DIR / "data" / "processed" / "vital_detectability_required_fields.csv",
        [
            "field_name",
            "required_for",
            "allowed_values",
            "action_if_unavailable",
        ],
    ),
    (
        "Appendix Table A21. Temporal red-priority tracking fields",
        BASE_DIR / "data" / "processed" / "vital_temporal_red_tracking_summary.csv",
        [
            "snapshot",
            "gene",
            "clinvar_id",
            "vital_score",
            "is_red",
            "global_ac",
            "popmax_ac",
            "qualifying_frequency_ac",
            "review_status",
            "inheritance_flag",
            "exit_reason_if_removed_from_red_queue",
        ],
    ),
    (
        "Appendix Table A22. Contextual overlay modules excluded from the core VITAL score",
        BASE_DIR / "data" / "processed" / "vital_contextual_overlay_modules.csv",
        [
            "overlay",
            "examples",
            "role",
            "score_policy",
        ],
    ),
    (
        "Appendix Table A23. Clinical decision projection for VITAL-red variants",
        BASE_DIR / "data" / "processed" / "vital_red_clinical_decision_projection.csv",
        [
            "gene",
            "clinvar_id",
            "current_public_label",
            "frequency_signal",
            "state_aware_frame",
            "if_generic_PLP_read_as_Mendelian_diagnosis",
            "if_state_aware_read_diagnosis",
            "one_label_two_incompatible_pathways",
            "clinical_projection_boundary",
        ],
    ),
    (
        "Appendix Table A24. ClinVar/ClinGen guideline-tension audit",
        BASE_DIR / "data" / "processed" / "vital_guideline_tension_audit.csv",
        [
            "gene",
            "clinvar_id",
            "current_public_label",
            "expert_panel_or_practice_guideline_review_present",
            "clinvar_state_terms_relevant",
            "guideline_tension",
            "why_this_is_not_a_ClinGen_disagreement_claim",
        ],
    ),
    (
        "Appendix Table A25. Maximum credible allele-frequency analysis",
        BASE_DIR / "data" / "processed" / "vital_max_credible_af_analysis.csv",
        [
            "gene",
            "clinvar_id",
            "scenario",
            "disease_prevalence",
            "allelic_contribution_fraction",
            "penetrance",
            "max_credible_af",
            "observed_popmax_af",
            "observed_to_max_credible_af_ratio",
            "constraint_call",
        ],
    ),
    (
        "Appendix Table A26. Dominant prevalence-constraint analysis",
        BASE_DIR / "data" / "processed" / "vital_prevalence_constraint_analysis.csv",
        [
            "gene",
            "clinvar_id",
            "scenario",
            "observed_popmax_af",
            "required_penetrance_for_observed_AF_to_fit_dominant_model",
            "expected_affected_prevalence_if_penetrance_10pct",
            "expected_affected_prevalence_if_penetrance_50pct",
        ],
    ),
    (
        "Appendix Table A27. Bayesian ACMG-style frequency tension",
        BASE_DIR / "data" / "processed" / "vital_bayesian_acmg_frequency_tension.csv",
        [
            "gene",
            "clinvar_id",
            "frequency_conflict_weight",
            "benign_frequency_likelihood_ratio_divisor",
            "pathogenic_evidence_scenario",
            "pathogenic_likelihood_ratio_multiplier",
            "posterior_probability_after_frequency_and_pathogenic_evidence",
        ],
    ),
    (
        "Appendix Table A28. Recessive carrier logic for TRDN",
        BASE_DIR / "data" / "processed" / "vital_recessive_carrier_logic.csv",
        [
            "gene",
            "clinvar_id",
            "frequency_source",
            "allele_frequency_q",
            "heterozygote_carrier_frequency_approx_2q",
            "homozygote_frequency_q_squared",
            "dominant_expected_prevalence_if_penetrance_10pct",
            "dominant_expected_prevalence_if_penetrance_50pct",
        ],
    ),
    (
        "Appendix Table A29. Allele-count reliability for red-priority cases",
        BASE_DIR / "data" / "processed" / "vital_allele_count_reliability.csv",
        [
            "gene",
            "clinvar_id",
            "gnomad_qualifying_allele_count",
            "poisson_relative_standard_error_approx",
            "approx_95pct_relative_margin",
            "binomial_global_af_standard_error",
        ],
    ),
    (
        "Appendix Table A30. Tiered match reconciliation summary",
        BASE_DIR / "data" / "processed" / "vital_tiered_match_reconciliation_summary.csv",
        [
            "reconciliation_tier",
            "variants",
            "percent_of_total",
            "interpretation",
        ],
    ),
    (
        "Appendix Table A31. Tiered match reconciliation evidence layers",
        BASE_DIR / "data" / "processed" / "vital_tiered_match_reconciliation_layers.csv",
        [
            "evidence_layer",
            "variants",
            "percent_of_total",
            "meaning",
        ],
    ),
    (
        "Appendix Table A32. Provenance credibility-filter sensitivity summary",
        BASE_DIR / "data" / "processed" / "vital_provenance_credibility_filter_summary.csv",
        [
            "metric",
            "value",
            "note",
        ],
    ),
    (
        "Appendix Table A33. Provenance credibility-filter excluded cases",
        BASE_DIR / "data" / "processed" / "vital_provenance_credibility_filter_excluded_cases.csv",
        [
            "gene",
            "clinvar_id",
            "review_strength",
            "max_frequency_signal",
            "qualifying_frequency_ac",
            "vital_score",
            "vital_red_flag",
        ],
    ),
    (
        "Appendix Table A34. Clinical action context summary for arrhythmia frequency-tension variants",
        BASE_DIR / "data" / "processed" / "vital_clinical_action_context_summary.csv",
        [
            "clinical_context",
            "n_frequency_tension",
            "percent_of_tension",
            "ci95_low_percent",
            "ci95_high_percent",
            "ac_supported_n",
            "weak_review_n",
            "red_n",
            "popmax_only_n",
        ],
    ),
    (
        "Appendix Table A35. Operational regime distribution for arrhythmia frequency-tension variants",
        BASE_DIR / "data" / "processed" / "vital_frequency_tension_regime_distribution.csv",
        [
            "regime",
            "n",
            "percent_of_tension",
            "ci95_low_percent",
            "ci95_high_percent",
            "ac_ge_20_n",
            "weak_review_n",
            "red_n",
            "popmax_only_n",
            "af_gt_1e_4_n",
        ],
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 8) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_inline(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("\\", "")
    return text


def add_runs_with_basic_markdown(paragraph, text: str) -> None:
    text = clean_inline(text)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_table_from_rows(document: Document, rows: list[list[str]], font_size: int = 8) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            set_cell_text(cell, value, bold=row_index == 0, font_size=font_size)
            if row_index == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph()


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = []
    for offset, line in enumerate(table_lines):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        is_alignment = all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)
        if offset == 1 and is_alignment:
            continue
        rows.append(cells)
    return rows, index


def add_figure(document: Document, image_name: str, caption: str, figure_number: int) -> int:
    image_path = BASE_DIR / "figures" / image_name
    if not image_path.exists():
        return figure_number
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.8))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {figure_number}. {caption}")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    document.add_paragraph()
    return figure_number + 1


def add_markdown_image(document: Document, line: str, figure_number: int) -> int:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return figure_number
    alt, relative_path = match.groups()
    image_path = (BASE_DIR / relative_path).resolve()
    if not image_path.exists():
        return figure_number
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.8))
    caption = alt or image_path.stem.replace("_", " ")
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(f"Figure {figure_number}. {caption}.")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    document.add_paragraph()
    return figure_number + 1


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(16, 24, 32)


def add_appendix_tables(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Embedded Key Analysis Tables", level=1)
    for title, path, columns in APPENDIX_TABLES:
        if not path.exists():
            continue
        document.add_heading(title, level=2)
        with path.open("r", encoding="utf-8", newline="") as handle:
            reader = csv.DictReader(handle, delimiter="\t" if path.suffix == ".tsv" else ",")
            rows = [columns]
            for record in reader:
                row = []
                for column in columns:
                    value = record.get(column, "")
                    row.append(format_table_value(value))
                rows.append(row)
        add_table_from_rows(document, rows, font_size=7)


def format_table_value(value: str) -> str:
    if value is None:
        return ""
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return value
    if numeric != numeric:
        return ""
    if abs(numeric) >= 1000:
        return f"{numeric:,.0f}" if numeric.is_integer() else f"{numeric:,.2f}"
    if abs(numeric) >= 10:
        return f"{numeric:.1f}" if not numeric.is_integer() else f"{numeric:.0f}"
    if abs(numeric) >= 1:
        return f"{numeric:.3f}".rstrip("0").rstrip(".")
    if numeric == 0:
        return "0"
    return f"{numeric:.3g}"


def build_document() -> None:
    document = Document()
    configure_document(document)
    lines = MANUSCRIPT_MD.read_text(encoding="utf-8").splitlines()
    figure_number = 1
    index = 0

    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()
        if not line:
            index += 1
            continue

        if line.lstrip().startswith("|"):
            rows, index = parse_markdown_table(lines, index)
            add_table_from_rows(document, rows, font_size=8)
            continue

        if line.startswith("!["):
            figure_number = add_markdown_image(document, line, figure_number)
            index += 1
            continue

        if line.startswith("# "):
            document.add_heading(clean_inline(line[2:]), level=0)
            index += 1
            continue
        if line.startswith("## "):
            document.add_heading(clean_inline(line[3:]), level=1)
            index += 1
            continue
        if line.startswith("### "):
            document.add_heading(clean_inline(line[4:]), level=2)
            for image_name, caption in FIGURE_INSERTIONS.get(line, []):
                figure_number = add_figure(document, image_name, caption, figure_number)
            index += 1
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs_with_basic_markdown(paragraph, line[2:])
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            add_runs_with_basic_markdown(paragraph, re.sub(r"^\d+\. ", "", line))
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].rstrip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("![")
                or next_line.lstrip().startswith("|")
                or next_line.startswith("- ")
                or re.match(r"^\d+\. ", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            index += 1
        paragraph = document.add_paragraph()
        add_runs_with_basic_markdown(paragraph, " ".join(paragraph_lines))

    document.save(OUTPUT_DOCX)
    print(f"Saved {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_document()
